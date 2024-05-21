from django.contrib import admin
from .models import Feedback, AdminReply, Profile, CheckInOut, InOut
from import_export import resources
from import_export.admin import ImportExportModelAdmin
from django.db.models import Sum, Q
from django import forms
from datetime import datetime, timedelta
from django.http import HttpResponse
from docx import Document
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# Define resources for import/export functionality
class FeedbackResource(resources.ModelResource):
    class Meta:
        model = Feedback

# Inline admin class for AdminReply
class AdminReplyInline(admin.StackedInline):
    model = AdminReply
    extra = 0

# Form for selecting date range and export format
class ReportForm(forms.Form):
    date_range = forms.ChoiceField(choices=[
        ('day', 'Day'),
        ('week', 'Week'),
        ('month', 'Month'),
        ('year', 'Year'),
    ], required=True)
    export_format = forms.ChoiceField(choices=[
        ('pdf', 'PDF'),
        ('docx', 'DOCX'),
    ], required=True)

# Unregister models if they are already registered
models_to_unregister = [Feedback, AdminReply, Profile, CheckInOut, InOut]
for model in models_to_unregister:
    try:
        admin.site.unregister(model)
    except admin.sites.NotRegistered:
        pass

# Admin class for Feedback with additional functionalities
@admin.register(Feedback)
class FeedbackAdmin(ImportExportModelAdmin):
    resource_class = FeedbackResource
    list_display = ('user', 'timestamp', 'feedback', 'read_status')
    list_filter = ('user', 'timestamp')
    search_fields = ('user__username', 'feedback')
    inlines = [AdminReplyInline]
    actions = ['mark_as_read', 'send_bulk_replies']

    def read_status(self, obj):
        return obj.read
    read_status.boolean = True
    read_status.short_description = 'Read'

    def mark_as_read(self, request, queryset):
        queryset.update(read=True)
        self.message_user(request, "Selected feedbacks marked as read.")
    mark_as_read.short_description = "Mark selected feedbacks as read"

    def send_bulk_replies(self, request, queryset):
        for feedback in queryset:
            AdminReply.objects.create(feedback=feedback, reply="Thank you for your feedback!")
        self.message_user(request, "Replies sent to selected feedbacks.")
    send_bulk_replies.short_description = "Send bulk replies to selected feedbacks"

# Admin class for AdminReply
@admin.register(AdminReply)
class AdminReplyAdmin(admin.ModelAdmin):
    list_display = ('feedback', 'timestamp', 'reply')

# Admin class for Profile
@admin.register(Profile)
class ProfileAdmin(admin.ModelAdmin):
    list_display = ('user', 'birth_date', 'gender', 'address', 'hometown', 'job_position')
    search_fields = ('user__username', 'user__first_name', 'user__last_name', 'job_position')

class CheckInOutAdmin(admin.ModelAdmin):
    list_display = ('plate_number', 'vehicle_type', 'check_in_time', 'check_out_time', 'parking_fee')
    readonly_fields = ('parking_fee',)
    change_list_template = "admin/mvp/checkinout/change_list.html"

    def changelist_view(self, request, extra_context=None):
        # Handle form submission
        if 'apply' in request.GET:
            form = ReportForm(request.GET)
            if form.is_valid():
                date_range = form.cleaned_data['date_range']
                export_format = form.cleaned_data['export_format']

                today = datetime.today()
                if date_range == 'day':
                    start_date = today - timedelta(days=1)
                elif date_range == 'week':
                    start_date = today - timedelta(weeks=1)
                elif date_range == 'month':
                    start_date = today - timedelta(days=30)
                elif date_range == 'year':
                    start_date = today - timedelta(days=365)

                queryset = CheckInOut.objects.filter(
                    Q(check_in_time__gte=start_date) | Q(check_out_time__gte=start_date)
                )

                total_parking_fee = queryset.aggregate(Sum('parking_fee'))['parking_fee__sum'] or 0

                if export_format == 'pdf':
                    return self.export_as_pdf(queryset, total_parking_fee, date_range)
                elif export_format == 'docx':
                    return self.export_as_docx(queryset, total_parking_fee, date_range)

        else:
            form = ReportForm()

        extra_context = extra_context or {}
        extra_context['form'] = form

        response = super().changelist_view(
            request,
            extra_context=extra_context,
        )
        try:
            qs = response.context_data['cl'].queryset
            total_parking_fee = qs.aggregate(Sum('parking_fee'))['parking_fee__sum']
            if total_parking_fee:
                total_parking_fee = f"{total_parking_fee:.10f}"
            response.context_data['total_parking_fee'] = total_parking_fee
        except (AttributeError, KeyError):
            response.context_data['total_parking_fee'] = 0
        return response

    def export_as_pdf(self, queryset, total_parking_fee, date_range):
        if total_parking_fee:
            total_parking_fee = f"{total_parking_fee:.10f}"
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='MyCentered', alignment=1, fontName='Times-Roman'))
        styles.add(ParagraphStyle(name='MyTitle', alignment=1, fontName='Times-Roman', fontSize=24, spaceAfter=20, leading=30, bold=True))
        styles.add(ParagraphStyle(name='MyNormal', fontName='Times-Roman', fontSize=9, leading=13.5))
        styles.add(ParagraphStyle(name='MySignature', fontName='Times-Roman', fontSize=12))
        styles.add(ParagraphStyle(name='MySignatureContent', fontName='Times-Roman', fontSize=10, italic=True))

        elements = []
        
        elements.append(Paragraph("<b>SOCIALIST REPUBLIC OF VIETNAM</b>", styles['MyCentered']))
        elements.append(Paragraph("<b>INDEPENDENCE - FREEDOM - HAPPINESS</b>", styles['MyCentered']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("---------- o0o ----------", styles['MyCentered']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("REVENUE REPORT", styles['MyTitle']))
        
        elements.append(Paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}", styles['MyNormal']))
        elements.append(Paragraph(f"Total Parking Fee: {total_parking_fee}", styles['MyNormal']))
        elements.append(Paragraph(f"Date Range: {date_range}", styles['MyNormal']))
        
        data = [['Plate Number', 'Vehicle Type', 'Check-In Time', 'Check-Out Time', 'Parking Fee']]
        for record in queryset:
            row = [
                record.plate_number,
                record.vehicle_type,
                record.check_in_time.strftime('%Y-%m-%d %H:%M'),
                record.check_out_time.strftime('%Y-%m-%d %H:%M') if record.check_out_time else 'N/A',
                str(record.parking_fee)
            ]
            data.append(row)
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(Spacer(1, 12))
        elements.append(table)
        elements.append(Spacer(1, 48))

        # Signature sections
        data = [['Accountant', 'Director']]
        data.append(['(Signature and full name)', '(Signature and full name)'])

        table = Table(data, colWidths=[250, 250])
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTNAME', (0, 1), (-1, 1), 'Times-Roman'),
            ('FONTSIZE', (0, 1), (-1, 1), 10),
            ('ITALIC', (0, 1), (-1, 1)),
        ]))

        elements.append(table)
        
        doc.build(elements)
        buffer.seek(0)
        return HttpResponse(buffer, content_type='application/pdf')

    def export_as_docx(self, queryset, total_parking_fee, date_range):
        if total_parking_fee:
            total_parking_fee = f"{total_parking_fee:.10f}"
        document = Document()
        
        # Set document styles
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(9)
        style.paragraph_format.line_spacing = 1.5
        
        title_style = document.styles.add_style('MyTitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(24)
        title_font.bold = True
        title_paragraph_format = title_style.paragraph_format
        title_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        signature_style = document.styles.add_style('MySignatureStyle', WD_STYLE_TYPE.PARAGRAPH)
        signature_font = signature_style.font
        signature_font.name = 'Times New Roman'
        signature_font.size = Pt(12)
        signature_paragraph_format = signature_style.paragraph_format
        signature_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        signature_content_style = document.styles.add_style('MySignatureContentStyle', WD_STYLE_TYPE.PARAGRAPH)
        signature_content_font = signature_content_style.font
        signature_content_font.name = 'Times New Roman'
        signature_content_font.size = Pt(10)
        signature_content_font.italic = True
        signature_content_paragraph_format = signature_content_style.paragraph_format
        signature_content_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run('SOCIALIST REPUBLIC OF VIETNAM\nINDEPENDENCE - FREEDOM - HAPPINESS\n')
        run.bold = True
        
        p = document.add_paragraph('---------- o0o ----------', style='Normal')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        document.add_paragraph('REVENUE REPORT', style='MyTitleStyle')
        
        document.add_paragraph(f'Date: {datetime.today().strftime("%Y-%m-%d")}')
        document.add_paragraph(f'Total Parking Fee: {total_parking_fee}')
        document.add_paragraph(f'Date Range: {date_range}')
        
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Plate Number'
        hdr_cells[1].text = 'Vehicle Type'
        hdr_cells[2].text = 'Check-In Time'
        hdr_cells[3].text = 'Check-Out Time'
        hdr_cells[4].text = 'Parking Fee'
        
        for record in queryset:
            row_cells = table.add_row().cells
            row_cells[0].text = record.plate_number
            row_cells[1].text = record.vehicle_type
            row_cells[2].text = record.check_in_time.strftime('%Y-%m-%d %H:%M')
            row_cells[3].text = record.check_out_time.strftime('%Y-%m-%d %H:%M') if record.check_out_time else 'N/A'
            row_cells[4].text = str(record.parking_fee)
        
        document.add_paragraph()
        
        table = document.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = 'Accountant'
        table.cell(0, 1).text = 'Director'
        table.cell(1, 0).text = '(Signature and full name)'
        table.cell(1, 1).text = '(Signature and full name)'

        table.cell(0, 0).paragraphs[0].style = 'MySignatureStyle'
        table.cell(0, 1).paragraphs[0].style = 'MySignatureStyle'
        table.cell(1, 0).paragraphs[0].style = 'MySignatureContentStyle'
        table.cell(1, 1).paragraphs[0].style = 'MySignatureContentStyle'
        
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# Register the model with the admin class
admin.site.register(CheckInOut, CheckInOutAdmin)

# Admin class for InOut
@admin.register(InOut)
class InOutAdmin(admin.ModelAdmin):
    list_display = ('user', 'check_type', 'date', 'start_time', 'end_time')
    list_filter = ('check_type', 'date')
    search_fields = ('user__username', 'date')